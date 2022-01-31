import os
import re
import docx
from PyPDF2 import PdfFileReader
from docxtpl import DocxTemplate
from django.utils import timezone, dateformat
from Versionize.settings import BASE_DIR
from main.models import Document, Adjustment


def generate_info_card(data):
    """
    Формирует информационно-удостоверяющий лист в формате .docx на основании полученных из формы даты,
    а также queryset'а по таблицам document-section-project.
    Необходимый формат входных данных:
    data = {
            'document_id': ID документа для которого формируется документ (int),
            'developed_by': Данные из формы (str),
            'checked_by': Данные из формы (str),
            'norm_control': Данные из формы (str),
            'approved_by': Данные из формы (str),
            'manager_position': Данные из формы (str),
            'manager_name': Данные из формы (str),
    }
    """
    # Собираем данные для документа
    date = dateformat.format(timezone.now(), 'd.m.Y')

    # Данные из формы
    developed_by = data['developed_by']
    checked_by = data['checked_by']
    norm_control = data['norm_control']
    approved_by = data['approved_by']
    manager_position = data['manager_position']
    manager_name = data['manager_name']

    # Данные из БД
    document_section_data = Document.objects.select_related('section').get(id=data['document_id'])
    section_name = document_section_data.section.name
    section_abbreviation = document_section_data.section.abbreviation
    project_code = document_section_data.section.project.code
    section_code = f'{project_code}-{section_abbreviation}'
    project_name = document_section_data.section.project.name
    doc_version = document_section_data.version
    doc_variation = document_section_data.variation
    filename = document_section_data.name
    md5 = document_section_data.md5
    company_name = document_section_data.section.company.name

    # Метаданные файла
    file_path = os.path.join(BASE_DIR, 'media', 'main_docs', filename)
    with open(file_path, 'rb') as doc:
        reader = PdfFileReader(doc)
        metadata = reader.getDocumentInfo()
        # Разбираем данные о дате изменения файла
        pattern = r':(.{4})(.{2})(.{2})(.{2})(.{2})(.{2})'
        date_info = re.search(pattern, metadata['/ModDate']).groups()
        year = date_info[0]
        month = date_info[1]
        day = date_info[2]
        day = date_info[2]
        hours = date_info[3]
        minutes = date_info[4]
        seconds = date_info[5]
    file_meta_date = f'{day}.{month}.{year}'
    file_meta_time = f'{hours}:{minutes}:{seconds}'
    file_size = f'{os.path.getsize(file_path)} байт'

    # Формируем ИУЛ
    template_name = 'information_certificate_template.docx'
    doc = DocxTemplate(os.path.join(BASE_DIR, 'media', 'templates', template_name))
    context = {
        'section_code': section_code,
        'section_name': section_name,
        'project_name': project_name,
        'version': doc_version,
        'variation': doc_variation,
        'filename': filename,
        'md5': md5,
        'file_meta_date': file_meta_date,
        'file_meta_time': file_meta_time,
        'file_size': file_size,
        'developed_by': developed_by,
        'checked_by': checked_by,
        'norm_control': norm_control,
        'approved_by': approved_by,
        'manager_position': manager_position,
        'company_name': company_name,
        'manager_name': manager_name,
        'date': date,
    }
    doc.render(context)
    doc.save(os.path.join(BASE_DIR, 'media', 'downloads', f'{filename}_ИУЛ.docx'))


def remove_row(table, row):
    tbl = table._tbl
    tr = row._tr
    tbl.remove(tr)


def generate_permission_card(data):
    TEMPLATE_NAME = 'permission_card_template.docx'
    TEMPLATE_PATH = os.path.join(BASE_DIR, 'media', 'templates', TEMPLATE_NAME)

    # Считываем таблицу из файла, получаем объект
    doc = docx.Document(TEMPLATE_PATH)
    table = doc.tables[0]

    adjustment_data = Adjustment.objects.select_related().filter(document_id=data['document_id'])
    # Первая строка с данными о корректировках
    row_num = 3
    # Заполняем таблицу данными из БД по корректировкам
    for adjustment in adjustment_data:
        variation = str(adjustment.document.variation)
        pages = adjustment.pages
        body = adjustment.body
        code = adjustment.code
        note = adjustment.note

        variation_cell = table.cell(row_num, 0)
        pages_cell = table.cell(row_num, 1)
        body_cell = table.cell(row_num, 2)
        code_cell = table.cell(row_num, 4)
        note_cell = table.cell(row_num, 5)

        variation_cell.text = variation
        pages_cell.text = pages
        body_cell.text = body
        code_cell.text = code
        note_cell.text = note

        variation_cell.paragraphs[0].style = doc.styles['calibri_light_center']
        pages_cell.paragraphs[0].style = doc.styles['calibri_light_center']
        body_cell.paragraphs[0].style = doc.styles['calibri_light_left']
        code_cell.paragraphs[0].style = doc.styles['calibri_light_center']
        note_cell.paragraphs[0].style = doc.styles['calibri_light_left']

        row_num += 1
        table.add_row()
        # Строка была добавлена без объединённых ячеек
        merge_cell_1 = table.cell(row_num, 2)
        merge_cell_2 = table.cell(row_num, 3)
        merged_cell = merge_cell_1.merge(merge_cell_2)

    remove_row(table, table.rows[-1])

    # Генерируем новое название файла и сохраняем
    filename = adjustment_data.first().document.name[:-5]
    new_filename = os.path.join(BASE_DIR, 'media', 'downloads', f'{filename}_разрешение.docx')
    doc.save(new_filename)

    # Заполняем теги в файле
    doc = DocxTemplate(new_filename)

    project_code = adjustment_data.first().section.project.code
    section_abbreviation = adjustment_data.first().section.abbreviation
    context = {
        'section_abbreviation': f'{project_code}-{section_abbreviation}',
        'project_name': adjustment_data.first().section.project.name,
        'date': dateformat.format(timezone.now(), 'd.m.y'),
        'chief_engineer': adjustment_data.first().section.project.admin.last_name,
        'company_name': adjustment_data.first().section.company.name,
        # Следующие данные получаем из формы
        'permission_number': data['permission_number'],
        'norm_control': data['norm_control'],
        'changes_by': data['changes_by'],
        'made_by': data['made_by'],
        'approved_by': data['approved_by'],
    }
    doc.render(context)
    doc.save(new_filename)
